#!/usr/bin/env python3
"""Generate a Writer-style brief DOCX from the live Ask Zia landing page content."""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parents[1] / "briefs" / "ask-zia-landing.docx"


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_para(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        return p

    def add_bullet(text: str):
        return doc.add_paragraph(text, style="List Bullet")

    # Meta
    doc.add_heading("Ask Zia — AI Powered Data Analytics Landing Page", level=1)
    add_para("Source: https://www.zoho.com/analytics/zia/")
    add_para("Page type: Conversational analytics / AI assistant product landing")
    add_para("Title tag: AI Powered Data Analytics Software | Ask Zia - Zoho Analytics")
    add_para(
        "Meta description: Zoho Analytics is augmented with AI capabilities through Zia, "
        "our AI-powered data analytical assistant. Use Ask Zia natural language processing, "
        "Zia Insights, and generative AI."
    )

    # PAGE 1 — Hero
    doc.add_page_break()
    doc.add_heading("PAGE 1 — Hero", level=2)
    add_para("Eyebrow / category")
    add_para("Conversational analytics", bold=True)

    add_para("H1")
    doc.add_heading("AI powered Data Analytics with Zia", level=1)

    add_para("Hero body")
    add_para(
        "Zoho Analytics is augmented with AI capabilities through Zia, our AI-powered data "
        "analytical assistant. Use Ask Zia's natural language processing to access critical "
        "business insights in seconds, Zia Insights for automated insights, and more."
    )

    add_para("Primary CTA")
    add_para("Sign up for free", bold=True)

    add_para("Hero visual note")
    add_para("[Image placeholder: Ask Zia conversational UI / chat interface mockup]")
    add_para("TODO: replace with final hero asset")

    # PAGE 2 — Help strip
    doc.add_page_break()
    doc.add_heading("PAGE 2 — Here is how we can help you", level=2)
    doc.add_heading("Here is how we can help you", level=2)

    add_para("Icon row / capability chips (6 items)")
    help_items = [
        ("Conversational analysis", "Natural language questions to KPIs and visualizations"),
        ("Build business metrics", "Anyone can get insights regardless of querying skills"),
        ("Smart interpretation", "Smart suggestions; best-suited visualizations"),
        ("Predictive analysis", "Forecast trends and plan ahead"),
        ("AI driven smart insights", "Automated narrative insights from your data"),
        ("Generative AI", "ChatGPT / GPT-4 integrated for analytics tasks"),
    ]
    for title, desc in help_items:
        add_para(title, bold=True)
        add_para(desc)
        add_para("[Icon placeholder]")

    # PAGE 3 — Ask Zia features
    doc.add_page_break()
    doc.add_heading("PAGE 3 — Ask Zia feature sections", level=2)

    doc.add_heading("Ask Zia questions, get insights", level=2)
    add_para(
        "Ask Zia is a conversational AI assistant within Zoho Analytics, and she's here to "
        "help you go from raw data to actionable insights in seconds. Start a conversation "
        "with Zia, ask her anything, get meaningful insights as KPIs and powerful "
        "visualizations, and quickly arrive at critical business decisions. Your insights "
        "are now just a chat away!"
    )
    add_para("[Image placeholder: Ask Zia chat to visualization]")
    add_para("TODO: replace with final asset")

    doc.add_heading("Build business metrics", level=2)
    add_para(
        "Ask Zia puts the power of data in the hands of people. From CTOs to support agents, "
        "anyone can get actionable insights in seconds, regardless of their querying skills. "
        "Whether it is sales, marketing, or finance, Zia understands all your data and gives "
        "you 360° visibility into your business."
    )
    add_para("[Image placeholder: metrics / KPI cards]")
    add_para("TODO: replace with final asset")

    doc.add_heading("Smart interpretation", level=2)
    add_para(
        "Ask Zia provides smart suggestions, eliminates ambiguities and typos to create the "
        "best-suited visualizations based on your data. You can directly add these KPIs to "
        "any of your existing dashboards or build a new one from scratch."
    )
    add_para("[Image placeholder: smart suggestions / chart selection]")
    add_para("TODO: replace with final asset")

    doc.add_heading("Predictive analysis", level=2)
    add_para(
        "Ask Zia anything about future trends, predict scenarios, and plan ahead of time with "
        "smart forecasting. Arrive at data-driven, smart decisions and future proof your "
        "business with powerful insights."
    )
    add_para("[Image placeholder: forecast / predictive chart]")
    add_para("TODO: replace with final asset")

    # PAGE 4 — Zia Insights
    doc.add_page_break()
    doc.add_heading("PAGE 4 — Zia Insights", level=2)
    doc.add_heading("Zia Insights", level=2)
    add_para(
        "Insights don't have to be a needle in the haystack anymore. Save time gathering "
        "insights, and focus on turning those insights into action with our all new Zia insights."
    )

    doc.add_heading("AI-driven automated Insights", level=2)
    add_para(
        "Zia, with her deep understanding of your data, can now provide actionable insights "
        "in the form of easily understandable narratives about important trends in your data. "
        "Leverage the power of AI and unearth hidden insights that would otherwise be "
        "impossible or time-consuming to find in a single click."
    )
    add_para("[Image placeholder: Zia Insights narrative panel]")
    add_para("TODO: replace with final asset")

    # PAGE 5 — Generative AI + Mobile
    doc.add_page_break()
    doc.add_heading("PAGE 5 — Generative AI and Analyze on the go", level=2)

    doc.add_heading("Generative AI", level=2)
    add_para(
        "Zia is integrated with OpenAI's ChatGPT, and you can experience the power of GPT-4 "
        "seamlessly from within Zoho Analytics. Be it finding a public dataset and importing "
        "it into Zoho Analytics, creating formulas and generating SQL queries by just asking "
        "in natural language, or getting synonym suggestions for table column names, Zia's "
        "generative AI capabilities can help you solve a myriad of analytics tasks in pretty "
        "quick time."
    )
    add_para("CTA / link")
    add_para("Learn more", bold=True)
    add_para("[Image placeholder: Generative AI / ChatGPT in Zoho Analytics]")
    add_para("TODO: replace with final asset")

    doc.add_heading("Analyze on the go", level=2)
    add_para(
        "Zia is now available on your mobile device as an immersive chat-like interface. "
        "Carry your insights wherever you go, analyze on the move, and get powerful, "
        "real-time insights at your fingertip – anytime, anywhere."
    )
    add_para("[Image placeholder: mobile Ask Zia chat UI]")
    add_para("TODO: replace with final asset")

    # PAGE 6 — Business apps
    doc.add_page_break()
    doc.add_heading("PAGE 6 — Zia understands your business functions", level=2)
    doc.add_heading("Zia understands your business functions", level=2)
    add_para(
        "Zia understands data from popular business apps along with the domain lingo to "
        "offer contextual interactions and insights for you"
    )

    add_para("App / integration logos")
    for app in [
        "Zoho SalesIQ",
        "Zoho CRM",
        "Zoho Books",
        "Zoho Analytics",
        "QuickBooks",
        "Google Ads",
        "Zoho Desk",
    ]:
        add_bullet(app)
    add_para("More integrations", bold=True)
    add_para("[Logo strip placeholder — TODO: replace with final brand marks]")

    # PAGE 7 — Testimonial
    doc.add_page_break()
    doc.add_heading("PAGE 7 — What our customer says", level=2)
    doc.add_heading("What our customer says", level=2)

    add_para("Testimonial quote")
    add_para(
        '"The Ask Zia and Zia Insights features are also cool. I can ask for a specific '
        "agent name for their quarterly or monthly performance and get to see if their sales "
        "numbers are trending up or down, and from which geographic area their sales are "
        "coming from. I can also dive into why they can't close deals in other areas and try "
        "to replicate the successful ones. From Zia insights, I get which agent is doing the "
        "most amount of business and who is doing the least, and I can give their manager an "
        "update along with actionable insights—like which specific salesperson has a better "
        "understanding of customers on the coast rather than the mountains. These insights "
        "give a clear picture of the performance of all our sales reps, and we had this "
        'discovery within minutes because of Zia Insights."'
    )

    add_para("John Sheldon", bold=True)
    add_para("Business Intelligence Manager, Renu Energy Solutions")
    add_para("Read more", bold=True)
    add_para("[Customer photo / logo placeholder — TODO]")

    # PAGE 8 — How-Tos
    doc.add_page_break()
    doc.add_heading("PAGE 8 — Solutions / How-Tos", level=2)
    doc.add_heading("Solutions / How-Tos", level=2)

    for title in [
        "Getting started with Ask Zia",
        "How to train Zia",
        "Zia with ChatGPT",
    ]:
        add_para(title, bold=True)
        add_para("Watch the video")
        add_para("[Video thumbnail / icon placeholder — TODO]")

    # PAGE 9 — FAQs
    doc.add_page_break()
    doc.add_heading("PAGE 9 — Frequently Asked Questions", level=2)
    doc.add_heading("Frequently Asked Questions", level=2)

    faqs = [
        (
            "How do I get started with Zia?",
            'In Zoho Analytics, click the "Ask Zia" button on the left navigation bar and '
            "ask your questions to Zia.",
        ),
        (
            "Is my data safe with Zia?",
            "Yes. More than 60 million users trust Zoho. We have a strong privacy policy and "
            "security infrastructure.",
        ),
        (
            "What languages is Zia available in?",
            "Right now Zia is available in English and Spanish (available in beta).",
        ),
        (
            "Where can I learn more about Zia?",
            "You can refer to our help documentation to learn more about Zia.",
        ),
    ]
    for question, answer in faqs:
        doc.add_heading(question, level=3)
        add_para(answer)

    # PAGE 10 — Closing CTA
    doc.add_page_break()
    doc.add_heading("PAGE 10 — Closing banner / CTA", level=2)
    doc.add_heading("Ready to explore AI-powered analytics with Zia?", level=2)
    add_para(
        "Start asking questions in natural language, uncover automated insights, and put "
        "generative AI to work inside Zoho Analytics."
    )
    add_para("Primary CTA")
    add_para("Sign up for free", bold=True)
    add_para("Secondary CTA (optional)")
    add_para("Learn more about Ask Zia", bold=True)

    # Implementation notes
    doc.add_page_break()
    doc.add_heading("Implementation notes (for Web Page Builder)", level=2)
    add_para("Sections required (in order)", bold=True)
    for line in [
        "1. Hero — Conversational analytics eyebrow + H1 + body + Sign up for free CTA + hero visual",
        "2. Help-you strip — 6 capability icons (conversational, metrics, smart interpretation, predictive, smart insights, generative AI)",
        "3. Feature zigzag / stacked blocks — Ask Zia questions; Build business metrics; Smart interpretation; Predictive analysis",
        "4. Zia Insights intro + AI-driven automated Insights block",
        "5. Generative AI block + Learn more CTA",
        "6. Analyze on the go (mobile) block",
        "7. Business apps / integrations logo strip + More integrations",
        "8. Customer testimonial (John Sheldon / Renu Energy Solutions)",
        "9. Solutions / How-Tos video cards (3)",
        "10. FAQ accordion (4 questions)",
        "11. Closing CTA banner — Sign up for free",
    ]:
        add_para(line)

    add_para(
        "Do not invent copy beyond this brief. Use placeholder images with TODO comments. "
        "Nav/footer are template slots only."
    )
    add_para("Live reference: https://www.zoho.com/analytics/zia/")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print(f"Size {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
