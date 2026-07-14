#!/usr/bin/env python3
"""Generate 3 differently themed Writer-format DOCX files for tool testing."""

from docx import Document
from docx.shared import Pt

OUT_DIR = r"c:\Users\jegap\Downloads"


def add(doc, text, blank=False):
    doc.add_paragraph(text)
    if blank:
        doc.add_paragraph("")


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


# ---------------------------------------------------------------------------
# DOC 1 — Landing page format (same structure as testing.docx)
# Topic: Marketing dashboard examples
# ---------------------------------------------------------------------------
def build_marketing_landing():
    doc = new_doc()

    add(doc, "Marketing dashboard examples - Landing page")
    add(doc, "Explore 50+ Marketing Dashboard Examples & Templates")
    add(
        doc,
        "Browse real-world marketing dashboard examples built for growth teams. "
        "From campaign ROI to channel attribution, find a template that fits your "
        "data and get started in minutes.",
    )
    add(doc, "Explore Examples", blank=True)

    add(doc, "Section reference")
    add(doc, "Marketing Dashboard Examples by Type")
    add(
        doc,
        "Explore ready-to-use marketing dashboard templates for every use case, "
        "from campaign performance to funnel conversion. Each type includes "
        "multiple examples you can customize.",
    )

    types = [
        (
            "Campaign Performance Dashboard",
            "Track how every campaign is performing across channels in one view. "
            "Spot which campaigns drive the most conversions and which ones burn "
            "budget without results.",
            [
                "Total spend",
                "Impressions",
                "CTR",
                "Conversions",
                "Cost per acquisition",
                "ROAS by campaign",
            ],
        ),
        (
            "Channel Attribution Dashboard",
            "See which channels contribute to pipeline and revenue. Compare "
            "first-touch, last-touch, and multi-touch models so your team "
            "allocates budget with confidence.",
            [
                "Revenue by channel",
                "Pipeline by channel",
                "Attribution model comparison",
                "Assisted conversions",
                "Cost per lead by channel",
                "Channel mix over time",
            ],
        ),
        (
            "Content Marketing Dashboard",
            "Measure how blogs, videos, and email nurture contribute to awareness "
            "and conversions. Identify which content types drive engaged traffic "
            "and qualified leads.",
            [
                "Organic sessions",
                "Content conversion rate",
                "Top landing pages",
                "Email open and click rates",
                "Time on page",
                "Leads from content",
            ],
        ),
        (
            "Paid Ads Dashboard",
            "Monitor Google Ads, Meta, LinkedIn, and other paid channels side by "
            "side. Find wasted spend fast and double down on high-performing ad sets.",
            [
                "Ad spend by platform",
                "CPC and CPM",
                "Click-through rate",
                "Conversion rate",
                "ROAS",
                "Quality score trends",
            ],
        ),
        (
            "SEO Performance Dashboard",
            "Track organic visibility, keyword rankings, and search traffic growth. "
            "Connect SEO activity to leads and revenue so leadership sees the real impact.",
            [
                "Organic traffic",
                "Keyword ranking changes",
                "Click share",
                "Backlink growth",
                "Organic conversions",
                "Top landing pages",
            ],
        ),
        (
            "Marketing Funnel Dashboard",
            "Follow prospects from awareness to closed-won. Spot where drop-offs "
            "happen and which stages need better messaging or faster follow-up.",
            [
                "Visitors to leads",
                "Lead to MQL",
                "MQL to SQL",
                "SQL to opportunity",
                "Opportunity to closed-won",
                "Funnel conversion by stage",
            ],
        ),
        (
            "Email Marketing Dashboard",
            "Understand newsletter and nurture performance without exporting from "
            "your ESP. Compare campaigns, segments, and send times in one place.",
            [
                "Open rate",
                "Click rate",
                "Unsubscribe rate",
                "Revenue attributed to email",
                "List growth",
                "Campaign comparison",
            ],
        ),
        (
            "Social Media Dashboard",
            "Pull engagement and follower metrics across social platforms into a "
            "single marketing view. Connect social activity to website visits and form fills.",
            [
                "Engagement rate",
                "Reach and impressions",
                "Follower growth",
                "Clicks to site",
                "Best performing posts",
                "Share of voice",
            ],
        ),
    ]
    for title, desc, kpis in types:
        add(doc, title)
        add(doc, desc)
        add(doc, "Key KPIs")
        for k in kpis:
            add(doc, k)

    add(doc, "")
    add(doc, "Section reference")
    add(doc, "Create these dashboards in one-click")
    add(
        doc,
        "Simply authenticate your marketing tools, and Zoho Analytics will "
        "automatically generate a set of professional dashboards with all the "
        "essential metrics and KPIs.",
    )
    add(doc, "Build Marketing Dashboards Now")
    add(
        doc,
        "★ Takes only a few minutes   ★ Automated data sync   ★ Completely customizable",
        blank=True,
    )

    add(doc, "Section reference")
    add(doc, "Connect popular marketing tools you already use")
    add(
        doc,
        "Zoho Analytics integrates with the platforms your marketing team relies "
        "on every day. Get these marketing dashboards generated prebuilt for you "
        "as soon as you connect your tools.",
    )
    integrations = [
        (
            "Google Analytics",
            "Pull website traffic, conversions, and audience data into Zoho Analytics "
            "without manual exports. The integration syncs automatically.",
        ),
        (
            "Google Ads",
            "Connect Google Ads spend, clicks, and conversions to measure ROAS alongside "
            "your other channels.",
        ),
        (
            "Meta Ads",
            "Bring Meta campaign performance into Zoho Analytics to compare paid social "
            "with search and email.",
        ),
        (
            "HubSpot Marketing",
            "Sync HubSpot contacts, forms, and campaign data to build marketing reports "
            "beyond native HubSpot views.",
        ),
        (
            "Mailchimp",
            "Import Mailchimp campaign metrics and list growth into live dashboards your "
            "whole team can use.",
        ),
        (
            "LinkedIn Ads",
            "Connect LinkedIn Ads to track B2B campaign spend, leads, and cost per result "
            "next to your other channels.",
        ),
    ]
    for name, desc in integrations:
        add(doc, "logo")
        add(doc, name)
        add(doc, desc)

    add(doc, "")
    add(doc, "Section reference")
    add(doc, "How to create marketing dashboards using Zoho Analytics")
    add(doc, "Connect your data once. Your marketing dashboards stay updated automatically.")
    add(doc, "STEP 1")
    add(doc, "Connect your marketing tools")
    add(
        doc,
        "Link Google Analytics, Google Ads, Meta, HubSpot, Mailchimp, or any other "
        "marketing tool you use. Data syncs automatically, so your dashboards always "
        "reflect current numbers.",
    )
    add(doc, "STEP 2")
    add(doc, "Access prebuilt marketing dashboards")
    add(
        doc,
        "Get instant access to ready-made dashboards for campaigns, channels, SEO, "
        "funnel stages, and more. Just connect your data and the dashboards populate "
        "on their own.",
    )
    add(doc, "STEP 3")
    add(doc, "Build custom dashboards your way")
    add(
        doc,
        "Drag and drop charts, KPI widgets, and tables to build dashboards that match "
        "exactly what your team tracks. Filter by channel, campaign, region, or time "
        "period without writing a single query.",
    )
    add(doc, "STEP 4")
    add(doc, "Share dashboards with your team securely")
    add(
        doc,
        "Share live dashboards with marketers, agencies, or executives. Control who can "
        "view, edit, or export, so the right people see the right data.",
        blank=True,
    )

    add(doc, "Ready to build your marketing dashboard?")
    add(
        doc,
        "Start with a prebuilt template or build from scratch. Connect your marketing "
        "stack, set up your key metrics, and share live dashboards with your team, "
        "all in one place.",
    )
    add(doc, "START FOR FREE", blank=True)

    add(doc, "FAQs")
    faqs = [
        (
            "What is a marketing dashboard?",
            "A marketing dashboard is a visual report that pulls campaign, website, and "
            "CRM data into one screen. It tracks metrics like spend, conversions, ROAS, "
            "and channel performance in real time, so you do not have to dig through "
            "spreadsheets to know where growth stands.",
        ),
        (
            "What should a marketing dashboard include?",
            "At minimum: spend, conversions, cost per acquisition, ROAS, and channel mix. "
            "Many teams also add funnel stage conversion, SEO traffic, and content "
            "performance depending on what they manage.",
        ),
        (
            "Can I build a marketing dashboard without knowing SQL or coding?",
            "Yes. Zoho Analytics lets you connect your marketing tools and build "
            "dashboards using drag-and-drop and Ask Zia, the AI agent. You do not need "
            "to write any queries or formulas to get started.",
        ),
        (
            "Can I share my marketing dashboard with an agency or client?",
            "Yes. You can share dashboards with specific users via email, embed them "
            "in portals, or generate a view-only link. You control permissions for each share.",
        ),
        (
            "Can multiple marketers use the same dashboard with filtered views?",
            "Yes. You can set up user filters so each teammate only sees their own "
            "campaigns and channels when they log in, while managers see the full team "
            "view on the same dashboard.",
        ),
    ]
    for q, a in faqs:
        add(doc, q)
        add(doc, a)

    path = f"{OUT_DIR}\\writer-test-marketing-dashboard-examples.docx"
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# DOC 2 — Long-form comparison guide (same structure as cloud analytics.docx)
# Topic: Embedded analytics tools
# ---------------------------------------------------------------------------
def build_embedded_guide():
    doc = new_doc()

    add(doc, "H1: Embedded Analytics Tools Guide for 2026: Comparison Matrix")
    add(
        doc,
        "Your product team wants customers to see usage, billing, and outcomes inside "
        "your app. Today that means exporting CSVs, sending spreadsheet reports, or "
        "pointing users at a separate BI tool they never open. By the time your customer "
        "success team finishes the monthly pack, the numbers are already stale.",
    )
    add(
        doc,
        "Embedded analytics tools fix this at the product layer. They let you place "
        "interactive dashboards and reports inside your own application, with your "
        "branding, your customer isolation, and developer-friendly APIs. The question "
        "is which platform fits how you build and sell.",
    )
    add(
        doc,
        "Below we've compared the most widely used embedded analytics platforms in "
        "2026: what each one does well, where each falls short, and how to decide which "
        "one belongs in your product.",
    )
    add(doc, "Key Takeaways")
    add(
        doc,
        "Embedded analytics has become table stakes for SaaS products that want "
        "retention and upsell. The real differences: white-label depth, multi-tenant "
        "security, SDK flexibility, pricing transparency, and how much UI you can "
        "customize without forking the vendor app.",
    )
    add(
        doc,
        "64% of SaaS buyers expect product analytics inside the tools they already "
        "pay for. Native dashboards reduce support tickets and increase weekly active usage.",
    )
    add(
        doc,
        "Platforms fit different product shapes. Some optimize for iframe embeds, "
        "others for JS/React components, and a few for full white-label portals.",
    )
    add(
        doc,
        "The best fit depends on: tenant isolation needs, branding depth, engineering "
        "capacity, and whether analytics is a paid add-on.",
    )
    add(
        doc,
        "Zoho Analytics includes white-label embedding, 500+ connectors, and built-in "
        "AI from accessible pricing without enterprise feature gates.",
    )
    add(
        doc,
        "Most tools on this list offer a free trial. Testing with a real customer "
        "workspace matters more than reading vendor decks.",
        blank=True,
    )

    add(doc, "Side bar")
    for item in [
        "What is embedded analytics?",
        "Why embed analytics in your product?",
        "Embedded analytics tools at a glance",
        "Top 5 Embedded Analytics tools in 2026",
        "How to choose the right embedded analytics tool",
        "Key features to look for in an embedded analytics tool",
        "Why Zoho Analytics works for embedded analytics",
        "FAQs",
    ]:
        add(doc, item)
    add(doc, "")

    add(doc, "What is embedded analytics?")
    add(
        doc,
        "Embedded analytics is the practice of placing reports, dashboards, and "
        "self-service charts inside another application so end users never leave your "
        "product. Instead of exporting data to a standalone BI tool, customers explore "
        "insights where they already work.",
    )
    add(doc, "An embedded analytics platform:")
    for x in [
        "Connects to product databases, warehouses, and SaaS sources.",
        "Renders interactive dashboards inside your UI via iframe, SDK, or portal.",
        "Applies your logo, colors, and domain so analytics feels native.",
        "Isolates each customer tenant so accounts never see each other's data.",
        "Exposes APIs for SSO, row-level security, and theme control.",
    ]:
        add(doc, x)
    add(
        doc,
        "Here is what embedded analytics looks like in practice. A project management "
        "SaaS was emailing PDF reports every Friday. Clients complained the reports "
        "were late and incomplete. After embedding live dashboards inside the product "
        "with white-label branding, weekly engagement rose and support tickets about "
        "reporting dropped sharply.",
    )
    add(doc, "Ignore", blank=True)

    add(doc, "Why embed analytics in your product?")
    add(
        doc,
        "Standalone BI worked when customers had analysts on staff. Most of your users "
        "do not. They want answers inside the product they already open every morning.",
    )
    add(
        doc,
        "Products that ship embedded analytics see faster time-to-insight for customers, "
        "clearer upsell paths for premium analytics tiers, and fewer manual report "
        "requests for CS and engineering. The financial case is straightforward when "
        "analytics becomes a feature customers renew for, not a side project your team "
        "supports forever.",
    )
    add(
        doc,
        "Embedded analytics replaces the export-and-email treadmill with live, branded "
        "insights. The question is not whether to embed. It is how much longer you can "
        "ship product without it.",
    )
    add(doc, "CTA")
    add(
        doc,
        "Teams that embed analytics stop rebuilding monthly PDF packs. See what that "
        "looks like in your own product.",
    )
    add(doc, "Try Zoho Analytics for free today", blank=True)

    add(doc, "Embedded analytics tools at a glance")
    add(
        doc,
        "Tools | Best For | Embed style | White-label | From ($/user/mo) | Ideal Size | Built-in ML?",
    )
    for row in [
        "Zoho Analytics | SaaS & agencies | Embed / portal | Full | $8 | SMB → Enterprise | Yes – Zia AI",
        "Looker Embedded | GCP-native products | API / embed | Partial–Full | Custom | Mid → Enterprise | Yes – Vertex AI",
        "Tableau Embedded | Enterprise viz heavy | JS API | Partial | Custom | Enterprise | Partial",
        "Sisense | OEM / complex models | Embed SDK | Full | Custom | Mid → Enterprise | Yes",
        "Metabase | Startups needing speed | Embed iframe | Limited | Free / paid | Startup → SMB | No",
    ]:
        add(doc, row)
    add(doc, "")

    add(doc, "The below are the best 5 embedded analytics tools in 2026")
    add(
        doc,
        "1.Zoho Analytics - Best for SaaS products and agencies that need white-label "
        "depth without enterprise gates",
    )
    add(doc, "2.Looker Embedded - Best for GCP-native product teams")
    add(doc, "3.Tableau Embedded - Best for visualization-heavy enterprise products")
    add(doc, "4.Sisense - Best for OEM and complex data modeling")
    add(
        doc,
        "5.Metabase - Best for early-stage products that need a fast embed",
        blank=True,
    )

    add(doc, "Zoho Analytics - Best embedded analytics tool for SaaS and agencies")
    add(
        doc,
        "Zoho Analytics is a full-stack analytics platform that covers data connection, "
        "preparation, visualization, AI analysis, and white-label embedding. Built for "
        "teams that do not want to stitch together four vendors to ship one "
        "customer-facing dashboard.",
    )
    add(
        doc,
        "What makes the base price unusual is the breadth included. You get 500+ "
        "connectors, Ask Zia AI, self-service dashboards, white-label embedding, and "
        "cloud, on-prem, or hybrid deployment without enterprise-only feature walls.",
    )
    add(doc, "What Zoho Analytics does differently")
    add(
        doc,
        "Most embedded platforms either force deep engineering for every customization, "
        "or simplify so much that multi-tenant security becomes an afterthought. Zoho "
        "Analytics supports both self-service builders and developer APIs, without "
        "locking you into a single cloud ecosystem.",
    )
    add(doc, "How it actually works in practice")
    add(doc, "Zoho Analytics keeps everything in one place:")
    for x in [
        "Connect product DBs, warehouses, billing tools, and CS systems",
        "Model tenant-aware datasets once",
        "Build dashboards with drag-and-drop or queries",
        "Embed under your domain with SSO and branding",
    ]:
        add(doc, x)
    add(doc, "Key strengths")
    add(doc, "500+ native connectors")
    add(
        doc,
        "Zoho Analytics connects to databases, SaaS apps, files, and warehouses. Data "
        "can refresh on a schedule without custom ETL for common sources.",
    )
    add(doc, "White-label embedding")
    add(
        doc,
        "Publish dashboards under your domain with logos, colors, and navigation that "
        "match your product. Customers experience analytics as part of your brand.",
    )
    add(doc, "AI assistant (Zia)")
    add(
        doc,
        "Zia answers natural language questions, surfaces trends, and helps "
        "non-technical users explore data without opening a ticket.",
    )
    add(doc, "Flexible deployment")
    add(
        doc,
        "Cloud, on-premise, and hybrid options help teams with residency or compliance "
        "requirements.",
    )
    add(doc, "CTA")
    add(
        doc,
        "White-label embedding, connectors, and AI in one plan — no surprise feature gates.",
    )
    add(doc, "View all features and pricing")
    add(doc, "Pros:")
    for x in [
        "Pricing starts at $8/user/month with connectors and AI included",
        "Full white-label options suited to SaaS and agency portals",
        "Strong connector coverage reduces custom integration work",
        "Usable by both product engineers and business users",
    ]:
        add(doc, x)
    add(doc, "Cons:")
    for x in [
        "Complex OEM use cases may need deeper API planning up front",
        "Teams new to BI may need a short onboarding period",
        "Some advanced customizations require documentation or support",
    ]:
        add(doc, x)
    add(
        doc,
        '"With Zoho Analytics, customer-facing dashboards finally feel like our product '
        'instead of a bolted-on report."',
    )
    add(doc, "Product lead, B2B SaaS")
    add(doc, "Pricing")
    add(
        doc,
        "Plans start at $8/user/month. A 15-day free trial is available. No credit card "
        "required.",
    )
    add(doc, "View pricing plans")
    add(doc, "Who should choose Zoho Analytics?")
    add(
        doc,
        "SaaS companies embedding analytics for customers. Agencies delivering branded "
        "client portals. Teams that need connectors, AI, and white-label without a "
        "six-figure commitment.",
        blank=True,
    )

    # Tool 2–5 (shorter, still matching template)
    tools = [
        (
            "Looker Embedded - Best for GCP-native product teams",
            "Looker Embedded leans on LookML for consistent metrics and deep Google Cloud "
            "ties. Strong for teams already on BigQuery. Trade-offs include LookML "
            "learning curve and opaque pricing for smaller products.",
            ["LookML semantic layer", "BigQuery-native workflows", "Solid embedding APIs", "Enterprise governance"],
            ["Excellent when BigQuery is the warehouse", "Metric consistency at scale", "Mature developer ecosystem"],
            ["Steep learning curve", "Pricing can jump quickly", "Less compelling outside Google Cloud"],
            "Custom pricing. Limited self-serve trial.",
            "Who should choose Google Looker Embedded?",
            "Enterprise product teams on GCP with dedicated analytics engineering.",
        ),
        (
            "Tableau Embedded - Best for visualization-heavy enterprise products",
            "Tableau Embedded is strong when visual polish and analyst familiarity matter "
            "most. Embedding and licensing complexity can slow SaaS go-to-market.",
            ["Rich visualization library", "Familiar analyst tooling", "JS embedding options"],
            ["Best-in-class charting for many enterprise buyers", "Large talent pool already knows Tableau"],
            [
                "Licensing can be hard to forecast for multi-tenant products",
                "White-label depth varies by deployment",
                "Heavier than needed for early SaaS embeds",
            ],
            "Custom / capacity-based. Contact sales.",
            "Who should choose Tableau Embedded?",
            "Large enterprises shipping analytics to sophisticated internal or partner audiences.",
        ),
        (
            "Sisense - Best for OEM and complex modeling",
            "Sisense targets OEM analytics with flexible embedding and stronger data "
            "modeling for complex products. Pricing and implementation are typically "
            "enterprise-shaped.",
            ["OEM-friendly embedding", "Strong data modeling", "Custom analytics apps"],
            ["Good for deeply productized analytics", "Flexible for engineering-led teams"],
            [
                "Implementation often needs services",
                "Cost and complexity beyond SMB",
                "Overkill for simple dashboard embeds",
            ],
            "Custom pricing.",
            "Who should choose Sisense?",
            "Mid-market and enterprise OEM products with complex analytical models.",
        ),
        (
            "Metabase - Best for speed at the early stage",
            "Metabase is popular with startups that need embedded charts quickly. "
            "White-label and enterprise isolation features are more limited than "
            "dedicated OEM platforms.",
            ["Fast to embed", "Open-source roots", "Simple question builder"],
            ["Low barrier for early MVPs", "Transparent open-source community"],
            [
                "Limited white-label depth for paid SaaS tiers",
                "Less enterprise governance out of the box",
                "May outgrow you as tenant count grows",
            ],
            "Free tier plus paid plans.",
            "Who should choose Metabase?",
            "Early-stage products validating embedded reporting before committing to an OEM stack.",
        ),
    ]
    for title, intro, strengths, pros, cons, price, who_h, who_b in tools:
        add(doc, title)
        add(doc, intro)
        add(doc, "Key strengths")
        for s in strengths:
            add(doc, s)
        add(doc, "Pros")
        for p in pros:
            add(doc, p)
        add(doc, "Cons")
        for c in cons:
            add(doc, c)
        add(doc, "Pricing")
        add(doc, price)
        add(doc, who_h)
        add(doc, who_b, blank=True)

    add(doc, "How to choose the right embedded analytics tool for your product")
    add(
        doc,
        "Most teams pick from a demo instead of a real customer tenant. Use this "
        "grounded approach.",
    )
    steps = [
        (
            "Step 1: Map where product and customer data lives",
            "List your application DB, warehouse, billing, CRM, and support tools. "
            "Confirm each source can feed the embed without fragile nightly CSV jobs.",
        ),
        (
            "Step 2: Decide how native the experience must feel",
            "Iframe-only may be fine for an admin console. A paid analytics module usually "
            "needs theme, domain, and navigation control.",
        ),
        (
            "Step 3: Specify tenant isolation and SSO",
            "Define per-customer data filters, roles, and identity. If this is fuzzy, "
            "every vendor will look fine until the first security review.",
        ),
        (
            "Step 4: Calculate total cost of ownership",
            "Add seat pricing, embed SKUs, support packages, and engineering time. The "
            "cheapest headline rate is rarely the cheapest ship path.",
        ),
        (
            "Step 5: Run a real customer workspace through the trial",
            "Connect messy production-like data, embed in a staging app, and have CS and "
            "a customer smile test the result.",
        ),
    ]
    for h, b in steps:
        add(doc, h)
        add(doc, b)
    add(doc, "CTA")
    add(doc, "You've got the framework. Now run it against a real customer workspace.")
    add(doc, "Zoho Analytics gives you 15 days to test embedding for free. No credit card needed.")
    add(doc, "Start testing with your data", blank=True)

    add(doc, "Key features to look for in an embedded analytics tool")
    add(doc, "Feature | What to look for | Why it matters")
    for row in [
        "Data Connectivity | 300+ connectors; DB + SaaS + warehouse | Fewer connectors means more custom pipelines",
        "White-label | Logo, colors, domain, hide vendor branding | Customers must feel they are still in your product",
        "Multi-tenant security | RLS, tenant filters, SSO | Non-negotiable for SaaS",
        "Embed options | iframe, JS SDK, portal | Match your engineering stack",
        "Self-service | Drag-and-drop + NLP for end users | Reduces CS report requests",
        "AI/ML | Anomaly detection, forecasts, Ask AI | Turns static embeds into insight",
        "Pricing transparency | Clear embed and seat costs | Avoid invoice surprises at scale",
    ]:
        add(doc, row)
    add(doc, "")

    add(doc, "Why Zoho Analytics works for embedded analytics?")
    add(
        doc,
        "Most embedded platforms price OEM features at enterprise rates. The affordable "
        "ones struggle with branding and tenant isolation. Zoho Analytics aims to deliver "
        "both: embeddable analytics with white-label options and accessible pricing.",
    )
    add(
        doc,
        "The connector library covers typical product stacks — application databases, "
        "Stripe, HubSpot, warehouses — so day-one dashboards are realistic. Zia handles "
        "natural language questions without a separate ML stack. For SaaS companies, "
        "white-label embedding with customer-level isolation means clients see your "
        "brand, not a third-party BI chrome.",
    )
    add(
        doc,
        "Start your 15-day free trial and see embedding with your own workspace. Get "
        "full access. Upgrade anytime at a transparent price. No credit card required.",
    )
    add(doc, "Start free trial", blank=True)

    add(doc, "The complete embedded analytics platform - Zoho Analytics")
    add(
        doc,
        "The embedded analytics space has consolidated. Most platforms handle basic "
        "charts. Differences show up in white-label depth, tenant security, connector "
        "coverage, AI maturity, and pricing clarity.",
    )
    add(
        doc,
        "Zoho Analytics consistently wins evaluations where product teams care about "
        "more than one of those at once. Looker Embedded is hard to beat on GCP. Sisense "
        "is strong for complex OEM. For growing SaaS and agency products that need "
        "branded analytics without a six-figure commitment, Zoho Analytics is a practical "
        "place to start the trial.",
    )
    add(doc, "CTA")
    add(
        doc,
        "22,000+ businesses already run analytics on Zoho Analytics. Takes under a day "
        "to get your first embedded dashboard live.",
    )
    add(doc, "Join us - It's free to start", blank=True)

    add(doc, "FAQs")
    faqs = [
        (
            "What is an embedded analytics tool?",
            "An embedded analytics tool lets you place interactive reports and dashboards "
            "inside your own application so customers explore data without leaving your product.",
        ),
        (
            "Is embedded analytics the same as white-label reporting?",
            "Not exactly. Embedded analytics focuses on delivery inside your app. "
            "White-label reporting focuses on branding. Many teams need both so dashboards "
            "live in-product and look native.",
        ),
        (
            "Can I isolate data for each customer account?",
            "Yes. Platforms like Zoho Analytics support role-based access and row-level / "
            "tenant-level security so each account only sees its own data.",
        ),
        (
            "Do I need engineers to embed dashboards?",
            "Lightweight embeds can ship quickly with iframes or drop-in components. "
            "Deeper SDK themes and SSO usually need light engineering work, which is still "
            "far less than building BI from scratch.",
        ),
        (
            "Which tools are best for SaaS products?",
            "For SaaS, prioritize white-label depth, tenant isolation, SSO, and transparent "
            "pricing. Zoho Analytics, Sisense, and Looker Embedded are common evaluation "
            "shortlists depending on cloud stack and budget.",
        ),
        (
            "Is there an embedded analytics tool with built-in AI?",
            "Yes. Zoho Analytics includes Zia for natural language questions and assisted "
            "analysis. Other platforms include AI through their broader cloud stacks with "
            "more configuration overhead.",
        ),
    ]
    for q, a in faqs:
        add(doc, q)
        add(doc, a)

    path = f"{OUT_DIR}\\writer-test-embedded-analytics-guide.docx"
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# DOC 3 — Agency / client dashboard landing (different page type)
# Topic: Client reporting portal for marketing agencies
# ---------------------------------------------------------------------------
def build_agency_landing():
    doc = new_doc()

    add(doc, "Client reporting portal - Landing page")
    add(doc, "Client reporting portal software for agencies that outgrew monthly PDFs")
    add(
        doc,
        "Stop rebuilding the same client report by hand. Zoho Analytics helps agencies "
        "deliver live, branded client portals so every account can check campaign numbers "
        "anytime — without another status call or late-night export.",
    )
    add(doc, "Book a free demo", blank=True)

    add(doc, "Section reference")
    add(doc, "Everything an agency needs to run client reporting from one place")
    add(
        doc,
        "You're managing ten clients, each on a different mix of ads, analytics, and CRM "
        "tools, each expecting a report with their name on it. Zoho Analytics is built "
        "for that load. Connect every client's stack, build the portal once, and let the "
        "numbers refresh themselves.",
    )

    add(doc, "Connect every client's marketing stack")
    add(
        doc,
        "Every client runs a different mix of tools. One lives in Google Ads, another is "
        "all Meta, a third wants Search Console and email in the same view. Zoho Analytics "
        "connects to all of it and pulls the numbers in for you.",
    )
    for x in [
        "✔ Connect 30+ marketing tools including Google Analytics, Google Ads, Search Console, Meta Ads, LinkedIn Ads, Mailchimp, and HubSpot",
        "✔ Blend data across platforms into one cross-channel view no single tool can give you",
        "✔ Data syncs on a schedule you set, so you're not exporting files the night before a review",
    ]:
        add(doc, x)

    add(doc, "Spin up a new client portal in minutes")
    add(
        doc,
        "A new client signs up and they want to see marketing analytics from the first "
        "week. Start from a template built around the metrics agencies already track, "
        "then shape it to the account.",
    )
    for x in [
        "✔ 100+ pre-built reports and marketing dashboards across every connected tool",
        "✔ Ready-made templates for campaign performance, channel ROI, paid spend, and audience metrics",
        "✔ Customize any template to match that client's goals and the KPIs they actually care about",
    ]:
        add(doc, x)

    add(doc, "White-label every portal with your agency brand")
    add(
        doc,
        "Clients should open a report that looks like your product, not a third-party "
        "BI tool. Apply your logo, colors, and domain so every portal reinforces your agency.",
    )
    for x in [
        "✔ Apply agency logo, color palette, and custom domain",
        "✔ Remove vendor branding so the experience feels native to your studio",
        "✔ Keep each client's portal visually consistent while isolating their data",
    ]:
        add(doc, x)

    add(doc, "Offer AI-powered insights clients can use without calling you")
    add(
        doc,
        "Zia answers plain-English questions on top of each dashboard, summarizes what "
        "changed, and flags anomalies — so clients get context without scheduling another call.",
    )
    for x in [
        "✔ Ask Zia questions in natural language on live client data",
        "✔ Auto-generated summary narratives on key reports",
        "✔ Anomaly alerts when a KPI moves unexpectedly",
    ]:
        add(doc, x)

    add(doc, "Automate the monthly report email")
    add(
        doc,
        "Even with a live portal, some stakeholders still want a scheduled PDF or email "
        "summary. Automate that delivery on each client's cadence.",
    )
    for x in [
        "✔ Schedule branded report emails per client",
        "✔ Export or attach the views they care about most",
        "✔ Free your team from rebuilding the same pack every month",
    ]:
        add(doc, x)
    add(doc, "")

    add(doc, "Section reference")
    add(doc, "How agencies launch a client reporting portal with Zoho Analytics")
    add(doc, "Connect once. Brand once. Scale every new retainer without starting over.")
    add(doc, "STEP 1")
    add(doc, "Connect the client's data sources")
    add(
        doc,
        "Authenticate Google Analytics, ads accounts, CRM, and email tools for that "
        "client. Sync schedules keep numbers current without nightly spreadsheet exports.",
    )
    add(doc, "STEP 2")
    add(doc, "Choose or customize a portal template")
    add(
        doc,
        "Start from agency-ready marketing dashboards, then rename metrics, filters, and "
        "layouts to match the retainer's KPIs.",
    )
    add(doc, "STEP 3")
    add(doc, "Apply white-label branding and access")
    add(
        doc,
        "Add your logo, colors, and domain. Invite client users with view-only or "
        "filtered access so each login shows only their account.",
    )
    add(doc, "STEP 4")
    add(doc, "Share, automate, and iterate")
    add(
        doc,
        "Publish the live portal, schedule email reports if needed, and refine views as "
        "the engagement grows — without rebuilding from zero.",
        blank=True,
    )

    add(doc, "Section reference")
    add(doc, "Trusted by agencies that report every week")
    add(
        doc,
        "Marketing and digital agencies use Zoho Analytics to replace manual reporting "
        "with live client portals.",
    )
    add(doc, "G2")
    add(doc, "4.4 / 5")
    add(doc, "Capterra")
    add(doc, "4.5 / 5")
    add(doc, "GetApp")
    add(doc, "4.4 / 5")
    add(doc, "Software Advice")
    add(doc, "4.5 / 5", blank=True)

    add(doc, "Hear from our happy customers")
    add(
        doc,
        '"We stopped spending Sunday nights rebuilding decks. Clients open their portal '
        'whenever they want and we show up to strategy calls with live numbers."',
    )
    add(doc, "Operations Lead, Growth Agency")
    add(
        doc,
        '"White-label dashboards finally match our brand. Clients think we built the '
        'reporting product ourselves."',
    )
    add(doc, "Founder, Performance Marketing Studio")
    add(
        doc,
        '"Zia answers the basic client questions so my analysts work on the hard '
        'optimization problems."',
    )
    add(doc, "Analytics Manager, Full-service Agency", blank=True)

    add(doc, "Ready to give every client a portal they'll actually open?")
    add(
        doc,
        "See how Zoho Analytics pulls all your client data into branded, live portals, "
        "automates the monthly report, and hands clients AI that feels like your own.",
    )
    add(doc, "Book a free demo")
    add(
        doc,
        "★ No commitment required   ★ 30-minute demo   ★ See it working with your data",
        blank=True,
    )

    add(doc, "FAQs")
    faqs = [
        (
            "What is a client reporting portal?",
            "A client reporting portal is a live, branded dashboard environment where "
            "agency clients can check campaign and channel performance anytime instead "
            "of waiting for a monthly PDF.",
        ),
        (
            "Can I white-label dashboards with my agency's branding?",
            "Yes. Zoho Analytics lets you apply your agency's logo, colors, and domain "
            "to every client portal, so clients experience reporting as part of your "
            "service rather than third-party software.",
        ),
        (
            "How many data sources can I connect?",
            "Zoho Analytics connects 500+ data sources, including 30+ marketing tools "
            "like Google Analytics, Google Ads, Meta, and LinkedIn, plus databases, "
            "warehouses, and spreadsheets. You can blend them into one cross-channel view.",
        ),
        (
            "Can clients access their own dashboards securely?",
            "Yes. Each client gets secure, role-based access to a live portal showing "
            "only their own data. They can filter and drill in without emailing you for "
            "every question.",
        ),
        (
            "Does it automate client reporting emails?",
            "Yes. Zoho Analytics can sync data on a schedule and email each client a "
            "branded report on their cadence, so the monthly pack goes out without your "
            "team rebuilding it by hand.",
        ),
        (
            "How does the AI help with client reporting?",
            "Zia, the built-in AI, answers plain-English questions, generates summaries "
            "on reports and dashboards, and flags why a metric changed so clients get "
            "context behind the numbers.",
        ),
        (
            "Is Zoho Analytics good for agencies managing many clients?",
            "Yes. You can run separate portals for each client, control access "
            "individually, and reuse templates to set up new accounts quickly so "
            "reporting scales as your roster grows.",
        ),
    ]
    for q, a in faqs:
        add(doc, q)
        add(doc, a)

    path = f"{OUT_DIR}\\writer-test-client-reporting-portal.docx"
    doc.save(path)
    return path


def main():
    paths = [
        build_marketing_landing(),
        build_embedded_guide(),
        build_agency_landing(),
    ]
    for p in paths:
        print(f"OK: {p}")


if __name__ == "__main__":
    main()
